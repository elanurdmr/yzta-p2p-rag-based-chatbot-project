"""
app/api/upload_routes.py — Dosya yükleme ve ChromaDB indeksleme.
PDF, DOCX, DOC ve TXT destekleniyor. Her format için ayrı parser var.
"""

import asyncio
import io
import logging
from pathlib import Path

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import JSONResponse
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

upload_router = APIRouter(prefix="/upload", tags=["upload"])

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

# separators sırası önemli — önce paragraf ayrımı, sonra satır, sonra cümle...
# boş string en sona geliyor çünkü o son çare — karakter karakter keser
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1500,
    chunk_overlap=200,
    separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
    length_function=len,
)

# free tier dakikada ~100 embedding isteği kaldırıyor
# 40'lık batch + 5 saniye bekleme ile güvende kalıyoruz
EMBED_BATCH_SIZE = 40
EMBED_BATCH_DELAY = 5


# ---------------------------------------------------------------------------
# Metin çıkarma — her format için ayrı fonksiyon, hepsi Document listesi döndürür
# ---------------------------------------------------------------------------

def _documents_from_pdf(filename: str, content: bytes) -> list[Document]:
    """Her sayfayı ayrı Document olarak çıkarır — sayfa numarası metadata'da korunuyor."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    documents: list[Document] = []
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        # boş sayfaları atla — genellikle görsel sayfalar oluyor
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": filename, "page": page_num},
                )
            )
    return documents


def _documents_from_docx(filename: str, content: bytes) -> list[Document]:
    """DOCX — python-docx ile paragraf paragraf çıkar, tek Document'a birleştir."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return [Document(page_content=text, metadata={"source": filename})]


def _documents_from_doc(filename: str, content: bytes) -> list[Document]:
    """Eski .doc formatı — mammoth ile ham metin çıkarıyoruz.
    Bazen tablolar ve grafikler kaybolabiliyor ama metin genellikle tam geliyor."""
    import mammoth

    with io.BytesIO(content) as f:
        result = mammoth.extract_raw_text(f)

    text = result.value
    if not text.strip():
        raise ValueError(
            ".doc dosyasından metin çıkarılamadı. "
            "Lütfen dosyayı .docx formatına dönüştürüp tekrar deneyin."
        )
    return [Document(page_content=text, metadata={"source": filename})]


def _documents_from_txt(filename: str, content: bytes) -> list[Document]:
    """TXT — UTF-8 decode, errors='ignore' ile bozuk karakterleri atla."""
    text = content.decode("utf-8", errors="ignore")
    return [Document(page_content=text, metadata={"source": filename})]


def _extract_documents(filename: str, content: bytes) -> list[Document]:
    """Uzantıya göre doğru parser'ı seç ve çalıştır."""
    ext = Path(filename).suffix.lower()
    dispatch = {
        ".pdf": _documents_from_pdf,
        ".docx": _documents_from_docx,
        ".doc": _documents_from_doc,
        ".txt": _documents_from_txt,
    }
    if ext not in dispatch:
        raise ValueError(
            f"Desteklenmeyen dosya uzantısı: {ext}. "
            f"İzin verilenler: {', '.join(sorted(dispatch))}"
        )
    return dispatch[ext](filename, content)


# ---------------------------------------------------------------------------
# ChromaDB yazma — retry + rate-limit koruması
# ---------------------------------------------------------------------------

async def _add_batch_with_retry(
    vector_store, batch: list[Document], max_retries: int = 4
) -> None:
    """429 gelirse üstel geri-çekilme ile tekrar dener.
    İlk bekleme 65 saniye — quota bir dakikada sıfırlanıyor."""
    loop = asyncio.get_running_loop()
    wait = 65
    for attempt in range(max_retries):
        try:
            await loop.run_in_executor(None, vector_store.add_documents, batch)
            return
        except Exception as exc:
            msg = str(exc)
            if any(k in msg for k in ("429", "Quota exceeded", "rate limit", "RESOURCE_EXHAUSTED")):
                if attempt < max_retries - 1:
                    logger.warning(
                        "Rate limit aşıldı — %ds bekleniyor (deneme %d/%d)...",
                        wait, attempt + 1, max_retries,
                    )
                    await asyncio.sleep(wait)
                    wait = min(wait * 2, 300)  # maksimum 5 dakika bekle
                else:
                    raise
            else:
                raise


async def _index_document(filename: str, content: bytes, thread_id: str = "") -> int:
    """Asıl indeksleme mantığı burada.
    Dosyayı parse et, chunk'lara böl, thread_id ekle, ChromaDB'ye yaz.

    thread_id neden önemli: her kullanıcının kendi oturumuna yüklediği belgeler
    başka oturumların sorgularına karışmamalı. metadata filtresiyle izole ediyoruz.
    """
    from ai.rag.chromaClient import document_vector_store

    loop = asyncio.get_running_loop()
    # parse işlemi blocking olduğu için executor'a taşıyoruz
    page_docs = await loop.run_in_executor(None, _extract_documents, filename, content)

    if not page_docs or not any(d.page_content.strip() for d in page_docs):
        raise ValueError("Dosyadan metin çıkarılamadı veya dosya boş.")

    # split_documents kullanmak önemli — split_text değil
    # çünkü split_documents orijinal metadata'yı (source, page) tüm chunk'lara taşıyor
    chunks = text_splitter.split_documents(page_docs)

    if thread_id:
        for chunk in chunks:
            chunk.metadata["thread_id"] = thread_id
    total = len(chunks)
    logger.info("'%s' → %d parça oluşturuldu, ChromaDB'ye yazılıyor...", filename, total)

    for i in range(0, total, EMBED_BATCH_SIZE):
        batch = chunks[i : i + EMBED_BATCH_SIZE]
        logger.info("Grup %d–%d/%d ChromaDB'ye yazılıyor...", i + 1, i + len(batch), total)
        await _add_batch_with_retry(document_vector_store, batch)
        if i + EMBED_BATCH_SIZE < total:
            logger.info("Grup tamamlandı, %ds bekleniyor...", EMBED_BATCH_DELAY)
            await asyncio.sleep(EMBED_BATCH_DELAY)

    return total


# ---------------------------------------------------------------------------
# Endpoint'ler
# ---------------------------------------------------------------------------

@upload_router.post("/document")
async def upload_document(
    file: UploadFile = File(...),
    thread_id: str = Form(""),
):
    """PDF, DOCX, DOC veya TXT dosyasını yükle ve ChromaDB'ye kaydet.

    thread_id gönderilirse belgeler bu oturuma özgü etiketlenir;
    sohbet sırasında yalnızca o oturuma ait belgeler sorgulanır.
    """
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Desteklenmeyen dosya türü: {ext}. "
                f"İzin verilenler: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
            ),
        )

    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Yüklenen dosya boş.")

    try:
        chunk_count = await _index_document(file.filename, content, thread_id)
        logger.info(
            "'%s' dosyası %d parçaya bölündü ve ChromaDB'ye kaydedildi. (thread: %s)",
            file.filename, chunk_count, thread_id or "global",
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    except Exception as exc:
        logger.error("Dosya indeksleme hatası: %s", exc)
        raise HTTPException(
            status_code=500,
            detail=f"Dosya işlenirken beklenmeyen hata oluştu: {exc}",
        )

    return JSONResponse(
        content={
            "message": f"'{file.filename}' başarıyla yüklendi ve {chunk_count} parçaya bölündü.",
            "filename": file.filename,
            "chunks": chunk_count,
        }
    )


@upload_router.get("/health")
async def upload_health():
    """Yükleme servisinin sağlık kontrolü."""
    return {
        "status": "ok",
        "supported_formats": sorted(ALLOWED_EXTENSIONS),
    }
